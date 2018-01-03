# To change this license header, choose License Headers in Project Properties.
# To change this template file, choose Tools | Templates
# and open the template in the editor.
__author__ = "joshsalazar"
__date__ = "$Dec 23, 2016 10:41:26 PM$"

from tkinter import *
from tkinter.filedialog import *
from tkinter.messagebox import *
from docx import Document

#DC = direccion de memoria
#Variables Globales---------------------------------------------------------------------------
path = ""   #Ruta del archivo
doc = ""    #String de todo el contenido del documento word
document = 0    #DC del doc. de word  
previewWindow = 0   #DC de la ventana donde s emuestra 
root = 0    #DC de la ventana principal
textWindow = 0  #DC del textWidget
cont = 0    #contador de los errores que tiene el doc
#tildesDiccionario = {}  #Dic que se crea con los archivos txt
#preposicionesDiccionario = {}   #Dic que se crea con los archivos txt
erroresDic = {}
listaErrores = []   #Lista de errores que contiene el doc
listaTags = []  #Lista de tags(cada tag representa un error)
listaPalabrasIE = []    #Lista de palabras interrogativas y exclamativas
posErrores = [] #Lista de las posiciones de los tags
IF = False  #bandera de interrogacion
#Funciones----------------------------------------------------------------------------------
def abrirVentanaDialogo():
    global path
    
    path =  askopenfilename()
    
def obtenerWord():
    global path,document
    
    document = Document(path)

def crearWord():
    d = Document()

    d.add_heading('Hamlet')
    d.add_heading('dramatis personae', 2)
    d.add_paragraph('Hamlet, the Prince of Denmark')

    d.save('hamlet.docx')

def ocultar(ventana):
    ventana.withdraw()
    
def mostrar(ventana):
    ventana.deiconify()

def leerArchivos(arch): #Func encargada de crear los diccionarios de los errores
    global erroresDic
    
    archivo = open(arch, "r", encoding="utf-8")
    lista = []  #Lista temporar para almacenar las lineas de cada txt
    for linea in archivo.readlines():
        lista.append(linea.split("/")) 
    for lista2 in lista:
        erroresDic[lista2[0]] = lista2[1][:-1]   #[:-1] Se le quita /n

def leerArchivos2(arch2):   #Funcion encargada de crear las listas de las palabras
    global listaPalabrasIE
    
    archivo = open(arch2, "r", encoding="utf-8")
    for linea in archivo.readlines():
        listaPalabrasIE.append(linea[:-1]) 
    

def marcarPosErrores(palabra):
    global doc,posErrores
    
    inicio = doc.find(palabra)
    fin = inicio+len(palabra)
    posErrores.append((inicio,fin))
    
    return inicio,fin   #Retorna un tupla de la pos inicial y final de la palabra
    
def nombre():   #Func encargada de crear los nombres de los tags
    global cont,listaTags
    
    resp =  "error"+str(cont)
    cont+=1
    listaTags.append(resp)
    
    return resp
def marcarPalabras(color):
    for tags in listaTags:
        textWindow.tag_config(tags, foreground=color)   #Recorre la lista de tags para marcarlos de color rojo en el doc
    
def detectarErrores():
    global document, doc, erroresDic, listaErrores, textWindow, listaTags
    
    for para in document.paragraphs:
        doc += para.text+"\n"   #Crea el string doc
        
    for llave in erroresDic: #Busca los errores en el diccionario
        if(llave in doc):   #Si esta el error, busca donde esta en el documento
            parOrdenado = marcarPosErrores(llave)
            textWindow.tag_add(nombre(), "1."+str(parOrdenado[0]),"1."+str(parOrdenado[1])) #crear el tag
            listaErrores.append(llave)    #lo agrega a la lista de errores
    marcarPalabras("red")
    

def corregirErrores():
    global listaErrores, posErrores, erroresDic, doc, textWindow
    
    for error in listaErrores:
        doc = doc.replace(error,erroresDic[error],1)    #Busca la solucion de cada error
        
    textWindow.delete(1.0, END)     #Elimina el string anterior
    textWindow.insert(INSERT,doc)   #Muestra el nuevo string corregido
    #marcarPalabras("green")
    
def detectarInicioOrac(listaSimbolos):
    global listaPalabrasIE, textWindow, doc

    inicio = 0
    palabra = ""
    linea = ""
    
    posFinalSimbolos = detectarOrac(listaSimbolos)    #Encuentra ls posicion donde hay un simbolo de la lista
    for cant in range(0, len(posFinalSimbolos)):    #Busca tantas veces como sea la cantidad de simbolos que debe completar
        linea = doc[inicio:posFinalSimbolos[cant]+(cant+1)]  #Agarra la posible linea donde debera completar el simbolo
        inicio = posFinalSimbolos[cant]+(cant+2)
        for palabra in listaPalabrasIE: #Busca si existe una palabra que indica que debe completar el simbolo
            if(linea.find(palabra) != -1):  #Si encuentra entonces lo arregla
                if(linea.find("?") != -1):
                    linea2 = linea.replace(palabra,"¿"+palabra,1)
                else:
                    linea2 = linea.replace(palabra,"¡"+palabra,1)
                doc = doc.replace(linea,linea2,1)   #Marca donde deberia ir algun simbolo
                break
    
    textWindow.delete(1.0, END)     #Elimina el string anterior
    textWindow.insert(INSERT,doc)   #Muestra el nuevo string corregido
    
def detectarOrac(listaSimbolo):
    global doc    

    posSimbolos = []
    num = -1 
    for simbolo in listaSimbolo:
        for cant in range(0,doc.count(simbolo)):
            num = doc.find(simbolo,num+1)
            posSimbolos.append(num)
    return posSimbolos


        
    
#Interfaz------------------------------------------------------------------------------------
def ventanaPrincipal():
    global root
    
    make = Button(root, text = "Make", relief = GROOVE, command = crearWord)
    search = Button(root, text = "Search", relief = GROOVE, command = abrirVentanaDialogo)
    preview = Button(root, text = "Preview", relief = GROOVE, command = lambda: ventanaPreview())

    make.place(height = 25, width = 60, x = 100, y = 270) #ancho,largo
    search.place(height = 25, width = 70, x = 160, y = 270)
    preview.place(height = 25, width = 75, x = 230, y = 270)

def ventanaPreview():
    global previewWindow,root, document,textWindow
    
    if(previewWindow != 0):
        ocultar(root)
        mostrar(previewWindow)
    elif(path != ""):
        ocultar(root)
        obtenerWord()
        
        previewWindow = Toplevel(height = 645, width = 600)
        
        textWindow = Text(previewWindow, state = NORMAL, wrap = WORD, relief = GROOVE)
        for para in document.paragraphs:
            textWindow.insert(INSERT, para.text)
            textWindow.insert(INSERT, "\n")
        #textWindow.config(state = DISABLED, relief = GROOVE)
        
        back = Button(previewWindow, text = "<", relief = GROOVE, command = lambda: mostrar(root) or ocultar(previewWindow))
        correct = Button(previewWindow, text = "Correct", relief = GROOVE, command = corregirErrores)
        save = Button(previewWindow, text = "Save", relief = GROOVE, command = lambda: detectarInicioOrac(["!","?"]))
        error = Button(previewWindow, text = "Error", relief = GROOVE, command = detectarErrores)
        
        back.place(height = 25, width = 30, x = 10, y = 10)
        textWindow.place(height = 570, width = 580, x = 10, y = 40)
        error.place(height = 25, width = 65, x = 410, y = 615)
        correct.place(height = 25, width = 65, x = 470, y = 615)
        save.place(height = 25, width = 50, x = 540, y = 615)
    else:
        showerror("Error", "Debe seleccionar algun archivo")
   
#Root---------------------------------------------------------------------------------------
root = Tk()
root.configure(height = 300, width = 310)
root.title("Bienvenido!")

leerArchivos("Tildes.txt")
leerArchivos("NormasPreposiciones.txt")
leerArchivos2("PalabrasIntExc.txt")
ventanaPrincipal()

root.mainloop()
